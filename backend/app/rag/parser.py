"""文档解析 — 多格式 → 纯文本，含表格结构保留"""
import os


def parse_document(file_path: str) -> tuple[str, dict]:
    """
    解析文档，返回 (纯文本, 元数据)
    元数据包含: total_pages, file_type
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return _parse_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return _parse_docx(file_path)
    elif ext == '.md':
        return _parse_text(file_path, 'markdown')
    elif ext == '.txt':
        return _parse_text(file_path, 'txt')
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def _parse_pdf(file_path: str) -> tuple[str, dict]:
    """PDF 解析 — 逐页提取文本 + 表格转 Markdown"""
    import fitz

    doc = fitz.open(file_path)
    total_pages = len(doc)

    pages_text = []
    for page_num, page in enumerate(doc, start=1):
        parts = [f"[第{page_num}页]"]

        # 1. 提取普通文本
        text = page.get_text()
        if text.strip():
            parts.append(text)

        # 2. 检测并提取表格 → Markdown 格式
        tables = _extract_tables(page)
        if tables:
            parts.append("\n" + tables + "\n")

        pages_text.append("\n".join(parts))

    doc.close()

    full_text = "\n\n".join(pages_text)
    metadata = {"total_pages": total_pages, "file_type": "pdf"}
    return full_text, metadata


def _extract_tables(page) -> str:
    """
    从 PDF 页检测表格，返回 Markdown 格式的表格文本。
    PyMuPDF 1.25+ 支持 find_tables()。
    """
    try:
        found = page.find_tables()
    except Exception:
        return ""

    if not found or not found.tables:
        return ""

    results = []
    for t_idx, table in enumerate(found.tables, start=1):
        try:
            rows = table.extract()
            if not rows or len(rows) < 2:
                continue
        except Exception:
            continue

        # 清洗：合并空行 / 去 None
        clean_rows = []
        for row in rows:
            cells = [str(c).replace("\n", " ").strip() if c else "" for c in row]
            # 跳过全空行
            if any(cells):
                clean_rows.append(cells)

        if len(clean_rows) < 2:
            continue

        # 构建 Markdown 表格
        md_lines = [f"\n**表 {t_idx}**\n"]
        col_count = max(len(r) for r in clean_rows)

        # 补齐列数不一致的行
        padded = [r + [""] * (col_count - len(r)) for r in clean_rows]

        # 表头
        md_lines.append("| " + " | ".join(padded[0]) + " |")
        # 分隔符
        md_lines.append("| " + " | ".join(["---"] * col_count) + " |")
        # 数据行
        for row in padded[1:]:
            md_lines.append("| " + " | ".join(row) + " |")

        results.append("\n".join(md_lines))

    return "\n\n".join(results)


def _parse_docx(file_path: str) -> tuple[str, dict]:
    """Word 文档解析 — 文本 + 表格转 Markdown"""
    from docx import Document

    doc = Document(file_path)
    parts = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag == "p":
            # 段落
            text = _get_docx_paragraph_text(element, doc)
            if text.strip():
                parts.append(text.strip())
        elif tag == "tbl":
            # 表格 → Markdown
            md = _docx_table_to_markdown(element, doc)
            if md:
                parts.append(md)

    full_text = "\n\n".join(parts)
    metadata = {"total_pages": 1, "file_type": "docx"}
    return full_text, metadata


def _get_docx_paragraph_text(element, doc) -> str:
    """提取 Word 段落纯文本"""
    from docx.text.paragraph import Paragraph
    try:
        p = Paragraph(element, doc)
        return p.text
    except Exception:
        return ""


def _docx_table_to_markdown(element, doc) -> str:
    """Word 表格 → Markdown"""
    from docx.table import Table
    try:
        table = Table(element, doc)
        rows = table.rows
        if len(rows) < 2:
            return ""

        md_lines = ["\n**表格**\n"]
        col_count = len(rows[0].cells)

        # 表头
        header = [rows[0].cells[i].text.replace("\n", " ").strip() for i in range(col_count)]
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("| " + " | ".join(["---"] * col_count) + " |")

        # 数据行
        for row in rows[1:]:
            cells = [row.cells[i].text.replace("\n", " ").strip() for i in range(len(row.cells))]
            # 补齐列数
            cells += [""] * (col_count - len(cells))
            md_lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(md_lines)
    except Exception:
        return ""


def _parse_text(file_path: str, file_type: str) -> tuple[str, dict]:
    """纯文本/Markdown 解析"""
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()

    metadata = {"total_pages": 1, "file_type": file_type}
    return text, metadata
